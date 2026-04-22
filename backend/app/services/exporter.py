from __future__ import annotations

from pathlib import Path

from loguru import logger

from app.services.document_store import ParsedDocument


def export_pdf(doc: ParsedDocument, annotations: dict, output_path: Path) -> Path:
    """Generate an annotated PDF from a ParsedDocument and its annotations."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError("reportlab not installed — run: uv add reportlab") from exc

    highlighted_indices = {h["sentence_idx"] for h in annotations.get("highlights", [])}
    snippets_by_idx: dict[int, list[dict]] = {}
    for sn in annotations.get("snippets", []):
        snippets_by_idx.setdefault(sn.get("sentence_idx", -1), []).append(sn)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocTitle", parent=styles["Title"],
        fontSize=20, spaceAfter=18, textColor=colors.HexColor("#18150F"),
    )
    normal_style = ParagraphStyle(
        "DocNormal", parent=styles["Normal"],
        fontSize=11, leading=18, spaceAfter=8, textColor=colors.HexColor("#18150F"),
    )
    highlighted_style = ParagraphStyle(
        "DocHighlighted", parent=normal_style,
        backColor=colors.HexColor("#FEFF9C"),
    )
    snippet_style = ParagraphStyle(
        "DocSnippet", parent=styles["Normal"],
        fontSize=9, leading=14, leftIndent=20, spaceAfter=6,
        textColor=colors.HexColor("#5C5650"),
        fontName="Helvetica-Oblique",
        borderPad=4,
    )

    doc_tpl = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    story = [Paragraph(doc.title, title_style), Spacer(1, 0.1 * inch)]

    for idx, sentence in enumerate(doc.sentences):
        style = highlighted_style if idx in highlighted_indices else normal_style
        story.append(Paragraph(sentence, style))
        for sn in snippets_by_idx.get(idx, []):
            note = f"<b>[Note on '{sn['term']}']:</b> {sn['explanation']}"
            story.append(Paragraph(note, snippet_style))

    doc_tpl.build(story)
    logger.info("event=pdf_exported doc_id={} path={}", doc.doc_id, output_path)
    return output_path


def export_docx(doc: ParsedDocument, annotations: dict, output_path: Path) -> Path:
    """Generate an annotated DOCX from a ParsedDocument and its annotations."""
    try:
        from docx import Document as DocxDocument
        from docx.enum.text import WD_COLOR_INDEX
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("python-docx not installed — run: uv add python-docx") from exc

    highlighted_indices = {h["sentence_idx"] for h in annotations.get("highlights", [])}
    snippets_by_idx: dict[int, list[dict]] = {}
    for sn in annotations.get("snippets", []):
        snippets_by_idx.setdefault(sn.get("sentence_idx", -1), []).append(sn)

    document = DocxDocument()

    document.add_heading(doc.title, level=1)

    for idx, sentence in enumerate(doc.sentences):
        para = document.add_paragraph()
        words = sentence.split()
        for i, word in enumerate(words):
            run = para.add_run(word + ("" if i == len(words) - 1 else " "))
            if idx in highlighted_indices:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        for sn in snippets_by_idx.get(idx, []):
            note_para = document.add_paragraph()
            note_para.style = "Normal"
            run = note_para.add_run(f"[Note on '{sn['term']}']: {sn['explanation']}")
            run.italic = True
            run.font.color.rgb = RGBColor(0x5C, 0x56, 0x50)
            run.font.size = Pt(9)

    document.save(str(output_path))
    logger.info("event=docx_exported doc_id={} path={}", doc.doc_id, output_path)
    return output_path
